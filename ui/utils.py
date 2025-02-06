# utils.py

import requests
import os
import tempfile
from pathlib import Path
import mimetypes
import numpy as np
import time
from datetime import datetime
import logging

# External libraries (replace with your own as needed)
import streamlit as st
from olympiabhub import OlympiaAPI

# from prompts import generate_meeting_analysis_prompts
from docx import Document
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

SERVER_URL = "http://0.0.0.0:8584"
TOKEN = os.getenv("OLYMPIA_TOKEN")
MODEL = os.getenv("OLYMPIA_MODEL")
# utils.py
from io import BytesIO
import re


def process_transcript_with_rag(transcript, timestamps):
    # Create a Document object from the transcript
    from llama_index.core import Document
    from llama_index.core import VectorStoreIndex, SimpleDirectoryReader

    documents = [Document(text=transcript)]

    # Create vector store index
    index = VectorStoreIndex.from_documents(documents, service_context=service_context)

    # Create query engine
    query_engine = index.as_query_engine()

    # Generate analysis using RAG
    analysis_prompt = "Analyze this transcript and provide a detailed summary"
    response = query_engine.query(analysis_prompt)

    return str(response)


def markdown_to_plain_text(markdown_text):
    # Remove headers (##)
    text = re.sub(r"#+\s*", "", markdown_text)

    # Convert bullet points to plain text
    text = re.sub(r"^\s*[-*]\s+", "• ", text, flags=re.MULTILINE)

    # Convert bold text
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)

    # Remove extra newlines
    text = re.sub(r"\n\s*\n", "\n\n", text)

    return text.strip()


def create_word_document(transcript, analyses_results):
    doc = Document()

    # Add analyses
    doc.add_heading("Analysis Results", 0)
    for analysis_type, content in analyses_results.items():
        doc.add_heading(analysis_type, level=1)
        plain_text = markdown_to_plain_text(content)
        doc.add_paragraph(plain_text)
        doc.add_page_break()
    # Add transcript

    doc.add_heading("Transcript", 0)
    doc.add_paragraph(transcript)
    doc.add_page_break()

    # Convert to bytes for download
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    return doc_bytes.getvalue()


def call_llm_API(PROMPT):

    api = OlympiaAPI(MODEL, TOKEN)

    # Format de messages façon OpenAI
    messages = [
        {
            "role": "system",
            "content": "Tu es un assistant specialise dans la synthese d'interview transcrit.",
        },
        {"role": "user", "content": PROMPT},
    ]

    # Sans proxy
    response = api.chat_completion(messages=messages, temperature=0.2, max_tokens=1500)
    return response["choices"][0]["message"]["content"]


def analize_text(input_text):
    PROMPT = "Fais moi un resume en bullet points de cet interview: /n " + input_text
    return call_llm_API(PROMPT)


def get_prompt(text_input, analysis_type):
    prompts = generate_meeting_analysis_prompts(text_input)
    return prompts.get(analysis_type, {}).get("prompt")


def get_all_analysis_types():
    return list(generate_meeting_analysis_prompts("").keys())


def ollama(message):
    return "Placeholder for ollama response."


def save_uploaded_file(uploaded_file):
    with tempfile.NamedTemporaryFile(
        delete=False, suffix=os.path.splitext(uploaded_file.name)[1]
    ) as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        return tmp_file.name


def get_file_type(filename):
    extension = os.path.splitext(filename)[1].lower()
    if extension == ".mp3":
        return "Audio"
    elif extension == ".mp4":
        return "Video"
    else:
        return None


def transcribe_audio(file_path):
    with open(file_path, "rb") as audio_file:
        response = requests.post(
            f"{SERVER_URL}/transcribe_audio",
            files={"file": (os.path.basename(file_path), audio_file, "audio/wav")},
            timeout=300,
        )
        if response.status_code == 200:
            return response.json()
        else:
            return {"error": f"Error: {response.status_code}"}


def transcribe_video(file_path):
    if not os.path.exists(file_path):
        return {"error": "File not found"}

    file_extension = os.path.splitext(file_path)[1].lower()
    mime_type = "video/mp4" if file_extension == ".mp4" else "audio/mpeg"
    filename = os.path.basename(file_path)

    with open(file_path, "rb") as media_file:
        files = {"file": (filename, media_file, mime_type)}
        response = requests.post(
            f"{SERVER_URL}/transcribe_video", files=files, timeout=300
        )
        if response.status_code == 200:
            return response.json()
        else:
            return {"error": f"Error: {response.status_code}"}


def create_accessible_timestamps(text, chunks, style_html):
    valid_chunks = []
    last_valid_end = 0

    for chunk in chunks:
        if not chunk.get("text", "").strip():
            continue

        timestamp = chunk["timestamp"]
        try:
            if isinstance(timestamp, list) and len(timestamp) >= 2:
                start_time = timestamp[0] * 60 + timestamp[1]
                end_time = last_valid_end + 0.1
                if len(timestamp) > 2:
                    end_time = timestamp[2] * 60 + timestamp[3]
                last_valid_end = end_time
            else:
                start_time = float(timestamp)
                end_time = last_valid_end + 0.1
                last_valid_end = end_time

            valid_chunks.append(
                {"text": chunk["text"].strip(), "timestamp": start_time}
            )
        except (TypeError, ValueError):
            continue

    result = [
        style_html,
        '<div class="transcript-container" role="region" aria-label="Transcript with timestamps">',
    ]
    last_end = 0

    for chunk in valid_chunks:
        chunk_text = chunk["text"]
        start = text.find(chunk_text, last_end)

        if start != -1:
            result.append(text[last_end:start])
            timestamp = chunk["timestamp"]

            minutes = int(timestamp) // 60
            seconds = int(timestamp) % 60
            time_display = f"{minutes:02d}:{seconds:02d}"

            result.append(
                f"""<span 
                class="timestamp-span"
                data-timestamp="{timestamp:.2f}"
                onclick="seekVideoAndPlay({timestamp:.2f})"
                role="button"
                tabindex="0"
                aria-label="Jump to {time_display}"
            ><span class="timestamp">[{time_display}]</span>{chunk_text}</span>"""
            )

            last_end = start + len(chunk_text)

    result.append(text[last_end:])
    result.append("</div>")

    result.append(
        """
        <script>
            function seekVideoAndPlay(timestamp) {
                const video = document.querySelector('video');
                if (!video) return;
                const seekTime = Math.max(0, timestamp);
                const seek = () => {
                    video.currentTime = seekTime;
                    video.play().catch(error => console.error("Playback failed:", error));
                };
                if (video.readyState >= 2) {
                    seek();
                } else {
                    video.addEventListener('loadeddata', function onLoad() {
                        seek();
                        video.removeEventListener('loadeddata', onLoad); 
                    });
                }
            }
            document.addEventListener('DOMContentLoaded', () => {
                document.querySelectorAll('.timestamp-span').forEach(span => {
                    span.addEventListener('keydown', e => {
                        if (e.key === 'Enter' || e.key === ' ') {
                            e.preventDefault();
                            const timestamp = parseFloat(span.dataset.timestamp);
                            seekVideoAndPlay(timestamp);
                        }
                    });
                });
            });
        </script>
    """
    )

    return "".join(result)


TIMESTAMP_STYLE = """
<style>
    .transcript-container {
        margin: 20px 0;
        line-height: 1.6;
        font-size: 16px;
    }
    .timestamp-span {
        display: inline;
        cursor: pointer;
        padding: 2px 4px;
        margin: 0 2px;
        border-radius: 3px;
        background-color: #e9ecef;
        transition: background-color 0.2s;
    }
    .timestamp-span:hover {
        background-color: #dee2e6;
    }
    .timestamp-span:focus {
        outline: 3px solid #0d6efd;
        outline-offset: 2px;
    }
    .timestamp {
        font-weight: bold;
        color: #0d6efd;
        margin-right: 4px;
        font-size: 14px;
    }
</style>
"""


def generate_meeting_analysis_prompts(text_input, language="FR"):
    system_prompt = """DIRECTIVES D'ANALYSE AUDIO/VIDÉO:
    
    TON RÔLE:
    Tu es un expert en analyse de contenus audiovisuels professionnels.
    
    FORMAT DE RÉPONSE:
    1. Résumé principal (5 lignes maximum) si demande
    2. Points essentiels par section
    3. Éléments factuels et mesurables
    4. Plan d'action concret
    
    RÈGLES DE STYLE:
    - Phrases courtes et précises (max 15 mots)
    - Verbes au présent uniquement
    - Vocabulaire professionnel
    - Format structuré avec puces
    - Données chiffrées en numériques
    
    INTERDICTIONS:
    - Pas de suppositions
    - Pas de conditionnel
    - Pas de familiarités
    - Pas d'interprétations personnelles
    - Pas de métadiscours
    - Pas d'interrogations
    - Pas de recommandations
    - Pas de formules de politesse
    - Pas d'avis subjectifs
    
    FORMATAGE:
    - Montants: XXX €
    - Pourcentages: XX%
    - Durées: HH:MM:SS
    - Listes: puces concises
    """

    base_prompts = {
        "compte_rendu": {
            "description": "Compte-rendu détaillé",
            "sections": [
                "RÉSUMÉ PRINCIPAL",
                "ORDRE DU JOUR",
                "DISCUSSIONS",
                "DÉCISIONS",
                "ACTIONS",
            ],
            "prompt": f"{system_prompt}\n\nGénération du compte-rendu:\n{text_input}",
        },
        "qui_dit_quoi": {
            "description": "Attribution des propos",
            "sections": [
                "INTERVENANTS",
                "CITATIONS",
                "PROPOSITIONS",
                "ENGAGEMENTS",
                "DÉSACCORDS",
            ],
            "prompt": f"{system_prompt}\n\nAnalyse des interventions par participant:\n{text_input}",
        },
        "questions_reponses": {
            "description": "Questions et réponses",
            "sections": [
                "QUESTIONS POSÉES",
                "RÉPONSES APPORTÉES",
                "POINTS EN SUSPENS",
                "CLARIFICATIONS",
                "SUIVIS",
            ],
            "prompt": f"{system_prompt}\n\nExtraction des questions et réponses:\n{text_input}",
        },
        "emotions": {
            "description": "Analyse émotionnelle",
            "sections": [
                "CLIMAT GÉNÉRAL",
                "TENSIONS",
                "CONSENSUS",
                "ENGAGEMENT",
                "DYNAMIQUE",
            ],
            "prompt": f"{system_prompt}\n\nAnalyse du climat émotionnel:\n{text_input}",
        },
        "statistiques": {
            "description": "Données statistiques",
            "sections": [
                "CHIFFRES CLÉS",
                "RATIOS",
                "ÉVOLUTIONS",
                "PROJECTIONS",
                "BENCHMARKS",
            ],
            "prompt": f"{system_prompt}\n\nExtraction des statistiques:\n{text_input}",
        },
        "mindmap": {
            "description": "Carte mentale",
            "sections": [
                "THÈME CENTRAL",
                "BRANCHES PRINCIPALES",
                "SOUS-THÈMES",
                "CONNEXIONS",
                "HIÉRARCHIE",
            ],
            "prompt": f"{system_prompt}\n\nCréation d'une carte mentale:\n{text_input}",
        },
        "points_focus": {
            "description": "Points d'attention majeurs",
            "sections": ["PRIORITÉS", "URGENCES", "IMPACTS", "CRITICITÉS", "VIGILANCE"],
            "prompt": f"{system_prompt}\n\nIdentification des points critiques:\n{text_input}",
        },
    }

    return base_prompts
